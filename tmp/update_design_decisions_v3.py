from copy import deepcopy
from pathlib import Path
import os

from docx import Document
from docx.text.paragraph import Paragraph


PATH = Path("World of Spirits - Updated v3.docx")
TEMP = Path("tmp/World of Spirits - Updated v3.editing.docx")


def add_after(paragraph, text, style="Normal"):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        if not child.tag.endswith("}pPr"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = style
    result.add_run(text)
    return result


def replace_range(doc, start_text, end_text, blocks, occurrence=0):
    paragraphs = doc.paragraphs
    starts = [i for i, p in enumerate(paragraphs) if p.text.strip() == start_text]
    start = starts[occurrence]
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip() == end_text)
    anchor = paragraphs[start]
    for paragraph in paragraphs[start + 1:end]:
        paragraph._element.getparent().remove(paragraph._element)
    for text, style in blocks:
        anchor = add_after(anchor, text, style)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


doc = Document(PATH)

replace_range(doc, "Game Level Design", "Game Lore", [
    ("Game Modes", "Heading 3"),
    ("World of Spirits has two main game modes: Story Mode and Infinity Mode.", "Normal"),
    ("Story Mode", "Heading 4"),
    ("Story Mode is a sequence of six separate stages. Each stage lasts ten minutes and ends with its elemental guardian boss. Defeating the boss completes the stage, unlocks the next stage, and returns the player to stage selection.", "Normal"),
    ("The stage order is Burning Plains, Frozen Wastes, Thunder Peaks, Poison Marsh, Shadow Realm, and Celestial Temple. Their guardians are the Fire Phoenix, Ice Yeti, Storm Dragon, Giant Scorpion, Necrotic Bat King, and Fallen Angel.", "Normal"),
    ("The player always begins a Story Mode stage with the spirit that matches that stage: Fire, Ice, Lightning, Poison, Necrotic, or Holy. Additional spirits can still be contracted through level-up choices during the run, up to the three-spirit limit.", "Normal"),
    ("Infinity Mode", "Heading 4"),
    ("Infinity Mode unlocks after all six Story Mode stages have been completed. It combines enemies and hazards from every plane into one continuous run. All six Story Mode bosses appear sequentially during the same run, with later encounters using stronger combinations and less recovery time. The run continues after the sixth boss for as long as the player survives.", "Normal"),
    ("Story Mode progression, stage unlocks, and Infinity Mode access are saved between sessions.", "Normal"),
])

replace_range(doc, "Starting Spirit System", "Leveling System", [
    ("In Story Mode, the starting spirit is determined by the selected stage and always matches that plane's element. Burning Plains starts with Fire, Frozen Wastes with Ice, Thunder Peaks with Lightning, Poison Marsh with Poison, Shadow Realm with Necrotic, and Celestial Temple with Holy.", "Normal"),
    ("Completing a stage unlocks the next stage. New spirits become available as starting options in non-story modes after their associated stage or summoning challenge has been completed.", "Normal"),
    ("During a run, the player can contract up to two additional spirits through level-up choices, for a maximum party size of three.", "Normal"),
])

# Fusion is intentionally removed until the core game is complete.
paragraphs = doc.paragraphs
fusion_index = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "Fusion System"), None)
if fusion_index is not None:
    progression_index = next(i for i in range(fusion_index + 1, len(paragraphs)) if paragraphs[i].text.strip() == "Progression System")
    for paragraph in paragraphs[fusion_index:progression_index]:
        remove_paragraph(paragraph)

replace_range(doc, "Run Rules", "Spirit Roster at a Glance", [
    ("A run supports up to three contracted spirits: one starting spirit and up to two spirits acquired through level-up choices.", "List Paragraph"),
    ("The player can manually rotate the spirit party. Rotation moves the next spirit into the main slot and shifts the previous main spirit into support.", "List Paragraph"),
    ("Spirit rotation has a one-second cooldown. Rotation input is ignored while this cooldown is active, and the UI must clearly show when rotation is ready.", "List Paragraph"),
    ("While the player is stationary, the main spirit becomes its elemental weapon and attacks automatically.", "List Paragraph"),
    ("While the player is moving, the main spirit returns to spirit form and casts support abilities. Both support spirits continue using abilities in either movement state.", "List Paragraph"),
    ("Each Story Mode stage lasts ten minutes and ends with a boss encounter.", "List Paragraph"),
    ("Spirit Rotation Buffs", "Heading 3"),
    ("Changing the main spirit grants a three-second elemental buff. Re-selecting the same buff refreshes its duration but does not stack its strength.", "Normal"),
    ("Fire - Blazing Resolve: increase all damage dealt by 20% for three seconds.", "List Paragraph"),
    ("Earth - Stoneguard: gain 25% damage reduction for three seconds.", "List Paragraph"),
    ("Wind - Tailwind: increase movement speed by 20% for three seconds.", "List Paragraph"),
    ("Water - Flow: ability cooldowns recover 25% faster for three seconds.", "List Paragraph"),
    ("Rotation buffs for Ice, Lightning, Poison, Necrotic, and Holy will be defined when those spirits enter production.", "List Paragraph"),
    ("Stationary Weapon Charging", "Heading 3"),
    ("After the player stops moving, the main spirit takes 0.25 seconds to transform into its weapon. This short transition prevents accidental weapon activation when the player only pauses briefly.", "Normal"),
    ("The weapon begins attacking immediately after transforming. Remaining stationary charges the weapon through two additional stages: Focused after one second and Empowered after two seconds.", "Normal"),
    ("Focused grants 15% weapon damage and 10% attack speed. Empowered grants 30% weapon damage and 20% attack speed. Charge does not increase beyond the Empowered stage.", "Normal"),
    ("Moving immediately returns the main spirit to support form, removes the weapon charge, and allows its support abilities to resume. Existing projectiles and persistent effects remain active until their normal duration ends.", "Normal"),
    ("The player, spirit, weapon, and HUD must visibly communicate transforming, Focused, Empowered, and rotation-cooldown states.", "Normal"),
    ("Status Effects", "Heading 3"),
    ("Burn: deals fire damage over time. Additional applications refresh duration and can stack up to a defined limit.", "List Paragraph"),
    ("Soaked: slows slightly and increases the damage of the next Lightning hit.", "List Paragraph"),
    ("Freeze: builds frost on a target; normal enemies freeze when the meter fills, while elites and bosses receive a shorter slow instead.", "List Paragraph"),
    ("Shock: deals lightning damage and can arc to nearby enemies.", "List Paragraph"),
    ("Poison: stacking damage over time that rewards repeated applications.", "List Paragraph"),
    ("Bleed: physical damage over time caused by spikes, blades, and crushing attacks.", "List Paragraph"),
    ("Pull, push, pin, and stun effects have reduced strength and duration against elites and bosses.", "List Paragraph"),
    ("Boss Resistance and Elemental Weakness", "Heading 3"),
    ("Bosses take 25% additional damage from their weakness and 25% less damage from their own element. These values are initial balance targets and can be tuned after playtesting.", "Normal"),
    ("Fire Phoenix: weak to Water; resistant to Fire.", "List Paragraph"),
    ("Ice Yeti: weak to Fire; resistant to Ice.", "List Paragraph"),
    ("Storm Dragon: weak to Earth; resistant to Lightning.", "List Paragraph"),
    ("Giant Scorpion: weak to Wind; resistant to Poison.", "List Paragraph"),
    ("Necrotic Bat King: weak to Holy; resistant to Necrotic.", "List Paragraph"),
    ("Fallen Angel: weak to Necrotic; resistant to Holy.", "List Paragraph"),
    ("Additional challenge bosses follow the same rule: Earth Golem is weak to Water, Water Leviathan is weak to Lightning, and Wind Roc is weak to Ice.", "List Paragraph"),
    ("Hard control never fully disables a boss. Freeze becomes a temporary slow, pin becomes heavy movement reduction, and stun becomes a brief interrupt with an internal resistance cooldown.", "Normal"),
])

# Remove the obsolete fusion open question.
for paragraph in list(doc.paragraphs):
    if "fusion" in paragraph.text.lower():
        remove_paragraph(paragraph)

replace_range(doc, "Fire Spirit", "Earth Spirit", [
    ("Shape", "Heading 3"), ("Phoenix", "Normal"),
    ("Weapon", "Heading 3"), ("Flame Bow", "Normal"),
    ("Description", "Heading 4"),
    ("A phoenix-forged bow that automatically fires burning feathers at nearby enemies. The Flame Bow rewards stationary weapon charging with rapid homing attacks and spreading Burn.", "Normal"),
    ("Weapon Levels", "Heading 4"),
    ("Lv1: Fires one homing feather at the nearest enemy.", "List Paragraph"),
    ("Lv2: Increases damage and projectile speed.", "List Paragraph"),
    ("Lv3: Fires two feathers and improves Burn buildup.", "List Paragraph"),
    ("Lv4: Burning targets explode when defeated, damaging nearby enemies.", "List Paragraph"),
    ("Lv5: Phoenix Bow - fires three piercing feathers and leaves a small fire patch after every third volley.", "List Paragraph"),
    ("Abilities", "Heading 3"),
    ("Ability 1: Fiery Feathers", "Heading 4"),
    ("Fires a fan of homing feathers that spreads Burn through groups.", "Normal"),
    ("Lv1: Fire 3 homing feathers in a fan.", "List Paragraph"),
    ("Lv2: Fire 5 feathers with improved tracking.", "List Paragraph"),
    ("Lv3: Feathers pierce one enemy and apply stronger Burn.", "List Paragraph"),
    ("Lv4: Feathers explode on their final hit.", "List Paragraph"),
    ("Lv5: Ashen Flock - fire 8 feathers; explosions leave short-lived fire patches.", "List Paragraph"),
    ("Ability 2: Fiery Talons", "Heading 4"),
    ("Leaves a burning trail behind the moving player, turning escape routes into damage zones.", "Normal"),
    ("Lv1: Leave a narrow fire trail that damages enemies over time.", "List Paragraph"),
    ("Lv2: Increase trail width and duration.", "List Paragraph"),
    ("Lv3: The trail applies Burn and deals more damage to already-burning enemies.", "List Paragraph"),
    ("Lv4: The trail periodically spreads flames toward nearby enemies.", "List Paragraph"),
    ("Lv5: Phoenix Footsteps - burning enemies defeated on the trail explode and extend its duration.", "List Paragraph"),
    ("Ability 3: Phoenix Dive", "Heading 4"),
    ("A flaming phoenix dives through a line of enemies and detonates at the end of its path.", "Normal"),
    ("Lv1: Perform one dive through the largest nearby enemy group.", "List Paragraph"),
    ("Lv2: Increase dive width and impact damage.", "List Paragraph"),
    ("Lv3: Leave a burning line along the dive path.", "List Paragraph"),
    ("Lv4: Perform a second crossing dive from another angle.", "List Paragraph"),
    ("Lv5: Rebirth Dive - the final impact creates a large fire zone and grants one revive per run at 30% health.", "List Paragraph"),
])

replace_range(doc, "Earth Spirit", "Water Spirit", [
    ("Shape", "Heading 3"), ("Golem", "Normal"),
    ("Weapon", "Heading 3"), ("Stone Hammer", "Normal"),
    ("Description", "Heading 4"),
    ("A massive stone hammer that sweeps around the player, crushing groups and rewarding long stationary charges with greater reach and impact.", "Normal"),
    ("Weapon Levels", "Heading 4"),
    ("Lv1: Sweep one hammer around the player.", "List Paragraph"),
    ("Lv2: Increase damage and rotation speed.", "List Paragraph"),
    ("Lv3: Increase reach and knockback.", "List Paragraph"),
    ("Lv4: The hammer creates a small shockwave on heavy hits.", "List Paragraph"),
    ("Lv5: Titan Hammer - summon a second hammer opposite the first; Empowered hits briefly stun normal enemies.", "List Paragraph"),
    ("Abilities", "Heading 3"),
    ("Ability 1: Quicksand Domain", "Heading 4"),
    ("Creates a slowing field around the player that controls crowds without copying Stone Spikes.", "Normal"),
    ("Lv1: Create a field that slows enemies by 20%.", "List Paragraph"),
    ("Lv2: Increase radius and slow strength.", "List Paragraph"),
    ("Lv3: The field deals damage over time.", "List Paragraph"),
    ("Lv4: Enemies are gradually pulled toward the center.", "List Paragraph"),
    ("Lv5: Sinking Kingdom - greatly increase the radius; normal enemies reaching the center are briefly immobilized and elites are heavily slowed.", "List Paragraph"),
    ("Ability 2: Boulder Throw", "Heading 4"),
    ("Throws a heavy boulder that bounces between enemy groups.", "Normal"),
    ("Lv1: Throw one boulder that bounces twice.", "List Paragraph"),
    ("Lv2: Add two bounces and increase damage.", "List Paragraph"),
    ("Lv3: Each bounce releases damaging rock fragments.", "List Paragraph"),
    ("Lv4: Direct hits stun normal enemies and briefly interrupt elites.", "List Paragraph"),
    ("Lv5: Continental Breaker - the final bounce causes a large explosion and splits into three smaller boulders.", "List Paragraph"),
    ("Ability 3: Stone Spikes", "Heading 4"),
    ("Stone spikes erupt beneath nearby enemy groups, pinning normal enemies and creating cracked ground.", "Normal"),
    ("Lv1: Summon 3 spikes near nearby enemies.", "List Paragraph"),
    ("Lv2: Summon 5 larger spikes; impaled enemies begin Bleeding.", "List Paragraph"),
    ("Lv3: Each spike sends a fault line toward another enemy, causing a smaller eruption.", "List Paragraph"),
    ("Lv4: Eruptions occur in two waves and leave damaging cracked ground.", "List Paragraph"),
    ("Lv5: Worldspine - summon a ring of massive spikes followed by a central eruption that launches enemies outward and deals bonus boss damage.", "List Paragraph"),
])

replace_range(doc, "Water Spirit", "Wind Spirit", [
    ("Shape", "Heading 3"), ("Leviathan", "Normal"),
    ("Weapon", "Heading 3"), ("Water Trident", "Normal"),
    ("Description", "Heading 4"),
    ("A leviathan-forged trident hurled through lines of enemies. It pierces targets on the outward path and returns to the player on a current.", "Normal"),
    ("Weapon Levels", "Heading 4"),
    ("Lv1: Throw one piercing trident.", "List Paragraph"),
    ("Lv2: Increase damage and travel speed.", "List Paragraph"),
    ("Lv3: The returning trident can damage enemies a second time and applies Soaked.", "List Paragraph"),
    ("Lv4: Throw a second trident in a nearby direction.", "List Paragraph"),
    ("Lv5: Leviathan's Reach - tridents create waves along their paths and return with increased size.", "List Paragraph"),
    ("Abilities", "Heading 3"),
    ("Ability 1: Tidal Wave", "Heading 4"),
    ("Sends waves outward to knock enemies away and create breathing room.", "Normal"),
    ("Lv1: Fire one wave in front of the player.", "List Paragraph"),
    ("Lv2: Fire a second wave behind the player.", "List Paragraph"),
    ("Lv3: Waves become wider and apply Soaked.", "List Paragraph"),
    ("Lv4: Fire waves in all four cardinal directions.", "List Paragraph"),
    ("Lv5: High Tide - release two consecutive rings of waves; the second wave deals bonus damage to Soaked enemies.", "List Paragraph"),
    ("Ability 2: Whirlpool", "Heading 4"),
    ("Summons whirlpools that pull enemies inward and group them for other attacks.", "Normal"),
    ("Lv1: Summon one whirlpool near an enemy group.", "List Paragraph"),
    ("Lv2: Increase radius and pull strength.", "List Paragraph"),
    ("Lv3: Summon two whirlpools.", "List Paragraph"),
    ("Lv4: Whirlpool deals damage over time and applies Soaked.", "List Paragraph"),
    ("Lv5: Maelstrom - the whirlpools slowly move together and merge into a larger damaging vortex.", "List Paragraph"),
    ("Ability 3: Rain Clouds", "Heading 4"),
    ("Rain clouds follow priority targets and continuously damage enemies beneath them.", "Normal"),
    ("Lv1: Summon one cloud over a nearby enemy.", "List Paragraph"),
    ("Lv2: Summon two clouds and improve target switching.", "List Paragraph"),
    ("Lv3: Increase rain damage and apply Soaked.", "List Paragraph"),
    ("Lv4: Clouds move faster and periodically release a damaging downpour.", "List Paragraph"),
    ("Lv5: Endless Monsoon - summon three larger clouds; overlapping rain zones create a lightning-vulnerable storm pool.", "List Paragraph"),
])

replace_range(doc, "Wind Spirit", "Ice Spirit", [
    ("Shape", "Heading 3"), ("Roc", "Normal"),
    ("Weapon", "Heading 3"), ("Chakrams", "Normal"),
    ("Description", "Heading 4"),
    ("The Roc throws wind-forged chakrams toward nearby enemies. They damage enemies while travelling outward, return to the player, and can strike the same enemy again on the return path.", "Normal"),
    ("Weapon Levels", "Heading 4"),
    ("Lv1: Throw one returning chakram.", "List Paragraph"),
    ("Lv2: Increase damage, speed, and return accuracy.", "List Paragraph"),
    ("Lv3: Chakrams pierce one additional enemy.", "List Paragraph"),
    ("Lv4: Throw two chakrams with a shorter cooldown.", "List Paragraph"),
    ("Lv5: Tempest Blades - throw three larger chakrams; catching each one releases a small wind burst.", "List Paragraph"),
    ("Abilities", "Heading 3"),
    ("Ability 1: Razor Wind", "Heading 4"),
    ("Fires radial wind blades that cut paths through surrounding enemies.", "Normal"),
    ("Lv1: Fire 2 blades in opposite directions.", "List Paragraph"),
    ("Lv2: Fire 4 blades.", "List Paragraph"),
    ("Lv3: Increase speed, range, and knockback.", "List Paragraph"),
    ("Lv4: Blades pierce three additional targets.", "List Paragraph"),
    ("Lv5: Thousand Cuts - fire 8 blades in two rotating waves; returning blades deal bonus damage.", "List Paragraph"),
    ("Ability 2: Tornado", "Heading 4"),
    ("Creates moving tornadoes that pull enemies inward and damage them repeatedly.", "Normal"),
    ("Lv1: Create one tornado in a ring five to nine units from the player.", "List Paragraph"),
    ("Lv2: Increase tornado radius and duration.", "List Paragraph"),
    ("Lv3: Increase pull strength and repeated damage.", "List Paragraph"),
    ("Lv4: Create two tornadoes moving in different directions.", "List Paragraph"),
    ("Lv5: Eye of the Storm - tornadoes grow while pulling enemies and release a final outward blast when they expire.", "List Paragraph"),
    ("Ability 3: Gale Barrier", "Heading 4"),
    ("A defensive wind pulse grants a temporary shield, damages nearby enemies, and blasts them away.", "Normal"),
    ("Lv1: Gain 10 shield and push nearby enemies away.", "List Paragraph"),
    ("Lv2: Increase barrier radius and shield to 16.", "List Paragraph"),
    ("Lv3: Destroy normal enemy projectiles touched by the pulse and grant 24 shield.", "List Paragraph"),
    ("Lv4: Release a second delayed gale pulse and grant 35 shield.", "List Paragraph"),
    ("Lv5: Sanctuary of Wind - the barrier persists briefly, deflects projectiles, and repeatedly pushes enemies away from the player.", "List Paragraph"),
])

replace_range(doc, "Fire Phoenix", "Earth Golem", [
    ("Encounter Structure", "Heading 3"),
    ("The current boss system supports phase changes at 66% and 33% health. Attacks are selected randomly without immediately repeating the previous attack, and each attack has a recovery period plus the boss's shared cooldown.", "Normal"),
    ("Phase 1 - The Hunt (100% to 66%)", "Heading 4"),
    ("Fire Dash and Feather Barrage teach the player to read lines and cones. Keep generous recovery windows so the player learns the patterns.", "Normal"),
    ("Phase 2 - Burning Sky (66% to 33%)", "Heading 4"),
    ("Flame Tornado becomes available, forcing the player away from marked positions and reducing safe space. Existing attacks may become slightly faster, but their warning shapes remain unchanged.", "Normal"),
    ("Phase 3 - Phoenix Storm (33% to 0%)", "Heading 4"),
    ("Meteor Rain becomes available. The Phoenix alternates movement pressure, projectiles, persistent tornadoes, and marked meteor impacts without repeating the same attack twice in a row.", "Normal"),
    ("Rebirth Phase", "Heading 4"),
    ("On its first death, the Phoenix stops attacking, becomes invulnerable for two seconds, triggers its rebirth presentation, and returns with 50% health. Rebirth is used only once. The boss remains in its most dangerous phase after returning.", "Normal"),
    ("Attack Telegraphs and Counterplay", "Heading 3"),
    ("Fire Dash: show a red line 10 units long and 1.5 units wide for 0.8 seconds. The line locks its direction before the dash. Safe response: step perpendicular to the line, then attack during recovery.", "List Paragraph"),
    ("Feather Barrage: the current attack waits 0.6 seconds before firing 7 feathers across a 70-degree fan. Add a visible fan-shaped warning, raised-wing animation, and sharp audio cue during this delay. Safe response: move behind the Phoenix or leave the fan.", "List Paragraph"),
    ("Flame Tornado: show a circular warning at the player's recorded position for 0.7 seconds, then create a moving tornado lasting 4 seconds. Safe response: leave the circle before it spawns and avoid being pushed into its path.", "List Paragraph"),
    ("Meteor Rain: mark 5 impact circles around the player's recorded position for 1 second while the meteors visibly fall. Safe response: keep moving through unmarked gaps and do not reverse into an old marker.", "List Paragraph"),
    ("Rebirth: remove normal attack warnings, darken the arena briefly, show the Phoenix collapsing into a flame core, then create a large harmless shockwave immediately before it returns. The two-second pause is a recovery and repositioning window.", "List Paragraph"),
    ("Elemental Rule", "Heading 3"),
    ("The Fire Phoenix takes 25% additional Water damage and 25% less Fire damage. Freeze and stun are converted into shorter slows or interrupts according to the boss-resistance rules.", "Normal"),
])

doc.save(TEMP)
os.replace(TEMP, PATH)
print(PATH)
