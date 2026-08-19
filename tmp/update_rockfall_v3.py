from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


PATH = Path("World of Spirits - Updated v3.docx")


def add_after(paragraph, text, style):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        if child.tag.endswith("}pPr"):
            continue
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = style
    result.add_run(text)
    return result


doc = Document(PATH)
paragraphs = doc.paragraphs
start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Ability 4:RockFall")
end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip() == "Water Spirit")
heading = paragraphs[start]
heading.text = "Ability 4: Rockfall"
for paragraph in paragraphs[start + 1:end]:
    paragraph._element.getparent().remove(paragraph._element)

blocks = [
    ("Description", "Normal"),
    ("Marks several locations near enemy groups before massive boulders fall from above. Each impact deals heavy area damage and leaves rubble that briefly blocks or redirects normal enemies.", "Normal"),
    ("Combat Role", "Normal"),
    ("A delayed bombardment ability that rewards positioning and controls enemy movement with temporary obstacles.", "Normal"),
    ("Upgrades", "Normal"),
    ("Lv1: Falling Stones - Mark 2 impact zones and drop a boulder on each after a short warning.", "List Paragraph"),
    ("Lv2: Heavy Rain - Drop 3 larger boulders with increased impact damage and radius.", "List Paragraph"),
    ("Lv3: Shattered Rock - Each boulder breaks into fragments that damage nearby enemies a second time.", "List Paragraph"),
    ("Lv4: Crushing Debris - Impacts leave rubble that slows enemies and forces normal enemies to move around it.", "List Paragraph"),
    ("Lv5: Mountainfall - Drop one enormous final boulder after the normal barrage. Its shockwave damages a wide area and briefly stuns elites.", "List Paragraph"),
]
anchor = heading
for text, style in blocks:
    anchor = add_after(anchor, text, style)

doc.save(PATH)
print(PATH)
