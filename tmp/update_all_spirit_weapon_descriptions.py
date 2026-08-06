from copy import deepcopy
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "World of Spirits - Updated.docx"
TEMP_PATH = ROOT / "tmp" / "World of Spirits - Updated.weapons.tmp.docx"


WEAPONS = {
    "Fire Spirit": (
        "Fire bow",
        "A phoenix-forged bow that automatically fires burning feathers at nearby enemies. "
        "The feathers home toward targets, and higher weapon levels improve damage, firing speed, "
        "projectile speed, and effective range.",
    ),
    "Earth Spirit": (
        "Stone hammer",
        "A massive stone hammer that continuously sweeps around the player while the Earth Spirit "
        "is in weapon form. Its broad orbit crushes groups of nearby enemies, with later levels "
        "improving damage, rotation speed, reach, and the number of targets struck.",
    ),
    "Water Spirit": (
        "Water trident",
        "A leviathan-forged trident hurled through lines of enemies with the force of the tide. "
        "Weapon upgrades increase its damage, throwing speed, attack rate, and targeting range.",
    ),
    "Wind Spirit": (
        "Chakrams",
        "The Roc throws wind-forged chakrams toward nearby enemies. Each chakram spins outward, "
        "travels up to seven units, then homes back to the player and disappears when caught. "
        "An enemy can be hit once on the outward trip and once again on the return trip. Damage, "
        "attack speed, projectile speed, projectile size, and multishot upgrades affect the weapon.",
    ),
    "Ice Spirit": (
        "Ice gauntlets",
        "Twin frozen gauntlets punch toward the nearest enemies. Their strikes can pierce clustered "
        "targets, knock enemies back, and briefly freeze anything they hit. Higher levels add a "
        "second fist and improve damage, reach, speed, and crowd control.",
    ),
    "Lightning Spirit": (
        "Lightning spear",
        "Design intent: a lightning-charged spear built for rapid piercing attacks. Its strikes will "
        "focus on speed, precision, and electrical damage that can arc from the primary target into "
        "nearby enemies.",
    ),
    "Poison Spirit": (
        "Poison daggers",
        "Design intent: venom-coated daggers thrown in quick succession. Each hit will apply poison "
        "or another weakening effect, rewarding sustained attacks against tougher enemies while "
        "still providing fast crowd pressure.",
    ),
    "Necrotic Spirit": (
        "Necrotic katana",
        "Design intent: a cursed katana that delivers fast sweeping cuts infused with necrotic "
        "energy. The weapon will specialize in finishing weakened enemies and spreading death-themed "
        "effects through tightly packed groups.",
    ),
    "Holy Spirit": (
        "Holy sword",
        "Design intent: a radiant sword that cleaves groups of corrupted enemies with wide arcs of "
        "holy energy. Its upgrades will emphasize area coverage, protection, and bonus damage against "
        "elite or corrupted targets.",
    ),
}


def insert_after(paragraph, text, style):
    clone = deepcopy(paragraph._p)
    for child in list(clone):
        if child.tag.endswith("}pPr"):
            continue
        clone.remove(child)
    paragraph._p.addnext(clone)
    result = paragraph._parent.add_paragraph()
    result._p.getparent().remove(result._p)
    result._p = clone
    result._element = clone
    result.style = style
    result.add_run(text)
    return result


doc = Document(DOC_PATH)

for spirit_name, (weapon_name, description) in WEAPONS.items():
    paragraphs = doc.paragraphs
    spirit_index = next(
        i for i, p in enumerate(paragraphs)
        if p.text.strip() == spirit_name and p.style.name == "Heading 2"
    )
    section_end = next(
        (i for i in range(spirit_index + 1, len(paragraphs))
         if paragraphs[i].style.name in ("Heading 1", "Heading 2")),
        len(paragraphs),
    )
    weapon_heading_index = next(
        i for i in range(spirit_index + 1, section_end)
        if paragraphs[i].text.strip() == "Weapon" and paragraphs[i].style.name == "Heading 3"
    )
    weapon_paragraph = paragraphs[weapon_heading_index + 1]
    weapon_paragraph.text = weapon_name

    paragraphs = doc.paragraphs
    weapon_index = next(i for i, p in enumerate(paragraphs) if p._p is weapon_paragraph._p)
    abilities_index = next(
        (i for i in range(weapon_index + 1, len(paragraphs))
         if paragraphs[i].text.strip() == "Abilities" and paragraphs[i].style.name == "Heading 3"),
        len(paragraphs),
    )
    description_heading = next(
        (paragraphs[i] for i in range(weapon_index + 1, abilities_index)
         if paragraphs[i].style.name == "Heading 4" and
         paragraphs[i].text.strip() in ("Description", "Weapon Behavior")),
        None,
    )

    if description_heading is None:
        description_heading = insert_after(weapon_paragraph, "Description", "Heading 4")
        insert_after(description_heading, description, "Normal")
    else:
        paragraphs = doc.paragraphs
        heading_index = next(i for i, p in enumerate(paragraphs) if p._p is description_heading._p)
        body = next(
            (paragraphs[i] for i in range(heading_index + 1, len(paragraphs))
             if paragraphs[i].text.strip()),
            None,
        )
        if body is None or body.style.name.startswith("Heading"):
            insert_after(description_heading, description, "Normal")
        else:
            body.text = description

doc.save(TEMP_PATH)
TEMP_PATH.replace(DOC_PATH)
print(DOC_PATH)
