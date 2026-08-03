from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "World of Spirits.docx"
OUTPUT = ROOT / "World of Spirits - Updated.docx"


def insert_before(anchor, element):
    anchor._p.addprevious(element._p if hasattr(element, "_p") else element._tbl)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_bullet(doc, text, bullet_num_pr):
    paragraph = doc.add_paragraph(style="List Paragraph")
    if bullet_num_pr is not None:
        paragraph._p.get_or_add_pPr().append(deepcopy(bullet_num_pr))
    paragraph.add_run(text)
    return paragraph


doc = Document(SOURCE)
bullet_source = next((p for p in doc.paragraphs if p.style.name == "List Paragraph"), None)
bullet_num_pr = (
    bullet_source._p.pPr.numPr
    if bullet_source is not None and bullet_source._p.pPr is not None
    else None
)

# Preserve the existing document's visual language while making new content easy to scan.
styles = doc.styles
if "Detail Note" not in [s.name for s in styles]:
    style = styles.add_style("Detail Note", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(55, 65, 81)
    style.paragraph_format.left_indent = Inches(0.2)
    style.paragraph_format.right_indent = Inches(0.2)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(8)

anchor = next(p for p in doc.paragraphs if p.text.strip() == "Spirits" and p.style.name == "Heading 1")

elements = []

heading = doc.add_paragraph("Current Gameplay Details", style="Heading 2")
elements.append(heading)

intro = doc.add_paragraph(style="Detail Note")
intro.add_run("Design snapshot: ").bold = True
intro.add_run(
    "The current game structure combines automatic combat, movement-based spirit behavior, "
    "and build choices across a ten-minute area run."
)
elements.append(intro)

h = doc.add_paragraph("Run Rules", style="Heading 3")
elements.append(h)
for text in (
    "A run supports up to three contracted spirits: one starting spirit and up to two spirits acquired through level-up choices.",
    "The active main spirit can be changed during play, allowing the player to switch weapon forms and combat roles.",
    "While the player is stationary, the main spirit becomes its elemental weapon and attacks automatically.",
    "While the player is moving, the main spirit casts abilities like a support spirit; support spirits continue attacking in both movement states.",
    "Each area lasts ten minutes and ends with a boss encounter.",
):
    elements.append(add_bullet(doc, text, bullet_num_pr))

h = doc.add_paragraph("Spirit Roster at a Glance", style="Heading 3")
elements.append(h)
table = doc.add_table(rows=1, cols=3)
table.autofit = False
table.columns[0].width = Inches(1.25)
table.columns[1].width = Inches(1.65)
table.columns[2].width = Inches(3.6)
headers = ("Spirit", "Form / Weapon", "Signature Abilities")
for idx, text in enumerate(headers):
    cell = table.rows[0].cells[idx]
    cell.text = text
    set_cell_shading(cell, "3F5F73")
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
roster = (
    ("Fire", "Phoenix / Fire bow", "Fiery Feathers, Fiery Talons, Phoenix Dive"),
    ("Earth", "Golem / Stone hammer", "Quicksand Domain, Boulder Throw, Stone Spikes"),
    ("Water", "Leviathan / Water trident", "Tidal Wave, Whirlpool, Rain Clouds"),
    ("Wind", "Roc / Chakrams", "Razor Wind, Tornado; third ability not yet defined"),
    ("Ice", "Yeti / Ice gauntlets", "Orbital Snowball, Avalanche, Ice Crystal"),
    ("Lightning", "Thunder Dragon / Lightning spear", "Lightning Strike, Chain Lightning Bolt, Thunder Roar"),
    ("Poison", "Scorpion / Poison daggers", "Toxic Glob, Venom Needles, Acid Spray"),
    ("Necrotic", "Bat / Necrotic katana", "Abilities not yet designed"),
    ("Holy", "Biblical Angel / Holy sword", "Healing, Shields, Light Beams (concept stage)"),
)
for row_idx, values in enumerate(roster, start=1):
    cells = table.add_row().cells
    for col_idx, value in enumerate(values):
        cells[col_idx].text = value
        cells[col_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cells[col_idx])
        if row_idx % 2 == 0:
            set_cell_shading(cells[col_idx], "EAF0F3")
        for paragraph in cells[col_idx].paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            for run in paragraph.runs:
                run.font.size = Pt(9)
elements.append(table)

h = doc.add_paragraph("Area Progression", style="Heading 3")
elements.append(h)
area_text = doc.add_paragraph(
    "The six-area sequence currently moves through Burning Plains, Frozen Wastes, Thunder Peaks, "
    "Poison Marsh, Shadow Realm, and Celestial Temple. Their final bosses are the Fire Phoenix, "
    "Ice Yeti, Storm Dragon, Giant Scorpion, Necrotic Bat King, and Fallen Angel."
)
elements.append(area_text)

h = doc.add_paragraph("Enemy Roles", style="Heading 3")
elements.append(h)
for text in (
    "Flying enemies pressure positioning and can approach over obstacles.",
    "Fast enemies pursue aggressively, while slow chargers telegraph heavier attacks.",
    "Ranged enemies force the player to keep moving and manage projectile space.",
    "Death-explosion enemies create temporary danger zones after defeat.",
):
    elements.append(add_bullet(doc, text, bullet_num_pr))

h = doc.add_paragraph("Open Design Questions", style="Heading 3")
elements.append(h)
for text in (
    "Define the long-term meta-progression system.",
    "Specify the exact ability levels, spirit pairings, and offer rules required to unlock fusion abilities.",
    "Complete the Wind Spirit's third ability and the Necrotic and Holy Spirit ability sets.",
    "Decide whether the Earth Golem and Wind Roc are challenge bosses, mini-bosses, or missing entries in the six-area sequence.",
):
    elements.append(add_bullet(doc, text, bullet_num_pr))

for element in elements:
    insert_before(anchor, element)

# Mark both data-table header rows so assistive technology and multi-page rendering
# can identify/repeat them correctly.
for data_table in doc.tables:
    tr_pr = data_table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")

# Remove the now-empty paragraphs left at the original append location.
body = doc._element.body
for child in list(body):
    if child.tag == qn("w:p") and not "".join(child.itertext()).strip() and child is not anchor._p:
        # Retain existing blank paragraphs; remove only appended empty artifacts is unnecessary.
        pass

doc.save(OUTPUT)
print(OUTPUT)
