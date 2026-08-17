from pathlib import Path
import sys
from docx import Document


def extract(path: Path) -> str:
    doc = Document(path)
    lines = [f"# {path.name}", ""]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = paragraph.style.name if paragraph.style else ""
            lines.append(f"[{style}] {text}")
    for index, table in enumerate(doc.tables, start=1):
        lines.extend(["", f"## TABLE {index}"])
        for row in table.rows:
            lines.append(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))
    return "\n".join(lines)


for value in sys.argv[1:]:
    path = Path(value)
    output = path.with_suffix(".extracted.txt")
    output.write_text(extract(path), encoding="utf-8")
    print(output)
