from copy import deepcopy
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "World of Spirits - Updated.docx"
TEMP_PATH = ROOT / "tmp" / "World of Spirits - Updated.tmp.docx"


def find_paragraph(doc, text, start=0):
    for index, paragraph in enumerate(doc.paragraphs[start:], start):
        if paragraph.text.strip() == text:
            return index, paragraph
    raise ValueError(f"Paragraph not found: {text}")


def insert_after(paragraph, text, style=None):
    new_paragraph = deepcopy(paragraph._p)
    for child in list(new_paragraph):
        if child.tag.endswith("}pPr"):
            continue
        new_paragraph.remove(child)
    paragraph._p.addnext(new_paragraph)
    result = paragraph._parent.add_paragraph()
    result._p.getparent().remove(result._p)
    result._p = new_paragraph
    result._element = new_paragraph
    if style is not None:
        result.style = style
    result.add_run(text)
    return result


def insert_bullet_after(paragraph, text, bullet_num_pr):
    item = insert_after(paragraph, text, "List Paragraph")
    if bullet_num_pr is not None:
        p_pr = item._p.get_or_add_pPr()
        existing = p_pr.numPr
        if existing is not None:
            p_pr.remove(existing)
        p_pr.append(deepcopy(bullet_num_pr))
    return item


doc = Document(DOC_PATH)
bullet_source = next((p for p in doc.paragraphs if p.style.name == "List Paragraph"), None)
bullet_num_pr = (
    bullet_source._p.pPr.numPr
    if bullet_source is not None and bullet_source._p.pPr is not None
    else None
)

# Upgrade selection is random and no longer gated by player level.
_, leveling = find_paragraph(doc, "Whenever the player levels up, they can choose an upgrade for one of their spirit abilities or gain a chance to obtain another spirit.")
leveling.text = (
    "Whenever the player levels up, the game randomly offers upgrades for spirit abilities, "
    "player stats, or new spirit contracts. Upgrade cards have no minimum player-level "
    "requirement; valid cards can appear from the beginning of a run. Spirit ownership, "
    "ability order, maximum ranks, and other functional prerequisites still apply."
)

# Update the summary table and retire the resolved Wind design question.
for table in doc.tables:
    if table.cell(0, 0).text.strip() == "Spirit" and table.cell(0, 1).text.strip() == "Form / Weapon":
        for row in table.rows[1:]:
            if row.cells[0].text.strip() == "Wind":
                row.cells[1].text = "Roc / Boomerang chakrams"
                row.cells[2].text = "Razor Wind, Tornado, Gale Barrier"
                break

_, open_question = find_paragraph(doc, "Complete the Wind Spirit's third ability and the Necrotic and Holy Spirit ability sets.")
open_question.text = "Complete the Necrotic ability set and finalize the Holy Spirit abilities."

# Add a concise implementation snapshot to Current Gameplay Details.
_, enemy_roles = find_paragraph(doc, "Enemy Roles")
implementation = insert_after(enemy_roles, "Recently Implemented Systems", "Heading 3")
last = implementation
for text in (
    "Upgrade offers are randomized without minimum player-level gates.",
    "Wind Tornadoes spawn in a random ring 5-9 units from the player, damage nearby enemies over time, and steadily pull enemies toward their center.",
    "Gale Barrier creates a visible expanding wind pulse, shields the player, and damages and pushes nearby enemies away.",
    "Wind Chakrams behave as boomerangs: they damage enemies while travelling outward, return to the player, and can damage the same enemy again on the return path.",
):
    last = insert_bullet_after(last, text, bullet_num_pr)

# Expand the dedicated Wind Spirit section.
wind_index, _ = find_paragraph(doc, "Wind Spirit")
_, chakrams = find_paragraph(doc, "Chakrams", wind_index)
weapon_heading = insert_after(chakrams, "Weapon Behavior", "Heading 4")
weapon_description = insert_after(
    weapon_heading,
    "The Roc throws wind-forged chakrams toward nearby enemies. Each chakram spins outward, "
    "travels up to seven units, then homes back to the player and disappears when caught. "
    "An enemy can be hit once on the outward trip and once again on the return trip. Damage, "
    "attack speed, projectile speed, projectile size, and multishot upgrades affect the weapon.",
    "Normal",
)
weapon_levels = insert_after(weapon_description, "Weapon Levels", "Heading 4")
last = weapon_levels
for text in (
    "Lv1: 8 damage, 0.90-second cooldown, 13-unit targeting range.",
    "Lv2: 11 damage with faster throws.",
    "Lv3: 14 damage, longer targeting range, and faster flight.",
    "Lv4: 18 damage with a substantially shorter cooldown.",
    "Lv5: 23 damage and maximum base throw speed.",
):
    last = insert_bullet_after(last, text, bullet_num_pr)

_, tornado_level_four = find_paragraph(doc, "Lv4 Two tornadoes.", wind_index)
tornado_notes = insert_after(tornado_level_four, "Implementation Notes", "Heading 4")
last = tornado_notes
for text in (
    "Tornadoes spawn randomly between five and nine units from the player rather than directly beside them.",
    "Each tornado lasts four seconds, deals damage every 0.5 seconds, and pulls enemies inward.",
    "The visible tornado remains at its authored size while its invisible pull radius scales independently.",
):
    last = insert_bullet_after(last, text, bullet_num_pr)

gale_heading = insert_after(last, "Ability 3: Gale Barrier", "Heading 4")
gale_description = insert_after(
    gale_heading,
    "A defensive wind pulse surrounds the player, grants a temporary shield, damages nearby "
    "enemies, and blasts them away. The pixel-art barrier expands to match the active area "
    "radius and then fades.",
    "Normal",
)
gale_upgrades = insert_after(gale_description, "Upgrades", "Normal")
last = gale_upgrades
for text in (
    "Lv1: Gain 10 shield and push nearby enemies away.",
    "Lv2: Larger barrier with 16 shield.",
    "Lv3: Stronger blast with 24 shield.",
    "Lv4: Twin gale blast with 35 shield.",
):
    last = insert_bullet_after(last, text, bullet_num_pr)

doc.save(TEMP_PATH)
TEMP_PATH.replace(DOC_PATH)
print(DOC_PATH)
