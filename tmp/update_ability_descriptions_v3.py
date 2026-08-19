from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


SOURCE = Path("World of Spirits - Updated v2.docx")
OUTPUT = Path("World of Spirits - Updated v3.docx")


def add_after(paragraph, text, style):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        if child.tag.endswith("}pPr"):
            continue
        new_p.remove(child)
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = style
    result.add_run(text)
    return result


def replace_section(doc, heading_text, next_heading_text, blocks):
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == heading_text)
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip() == next_heading_text)
    anchor = paragraphs[start]
    for paragraph in paragraphs[start + 1:end]:
        paragraph._element.getparent().remove(paragraph._element)
    for text, style in blocks:
        anchor = add_after(anchor, text, style)


doc = Document(SOURCE)

replace_section(
    doc,
    "Ability 2: Avalanche",
    "Ability 3: Ice Crystal",
    [
        ("Description", "Normal"),
        ("Launches a rolling snowball in the chosen direction. The snowball grows as it travels, damaging enemies and dragging smaller enemies along with it. At the end of its path, or when it strikes a solid obstacle, it explodes and fires ice shards in every direction.", "Normal"),
        ("Combat Role", "Normal"),
        ("A moving crowd-control projectile that gathers enemies into one place before finishing with a radial burst.", "Normal"),
        ("Upgrades", "Normal"),
        ("Lv1: Rolling Snowball - Launch a snowball that grows while rolling, pulls normal enemies with it, and explodes into 4 ice shards.", "List Paragraph"),
        ("Lv2: Gathering Momentum - The snowball grows faster, travels farther, deals more impact damage, and releases 6 ice shards.", "List Paragraph"),
        ("Lv3: Crushing Drift - Increases the pull radius and maximum size. Enemies carried by the snowball take repeated damage while being dragged.", "List Paragraph"),
        ("Lv4: Frozen Wake - The snowball leaves a freezing trail that slows enemies. Its explosion releases 8 shards that can pierce one target.", "List Paragraph"),
        ("Lv5: Glacial Catastrophe - Greatly increases the final explosion radius, freezes surviving enemies, and launches 12 high-damage ice shards.", "List Paragraph"),
    ],
)

replace_section(
    doc,
    "Ability 3: Ice Crystal",
    "Lightning Spirit",
    [
        ("Description", "Normal"),
        ("Creates ice crystals in a defensive ring around the player. Each crystal grows for a short time, damages nearby enemies, and shatters when an enemy comes close or when its duration ends. The shattering blast briefly freezes affected enemies.", "Normal"),
        ("Combat Role", "Normal"),
        ("A defensive trap ability that protects the player's immediate space and punishes enemies that push too close.", "Normal"),
        ("Upgrades", "Normal"),
        ("Lv1: Crystal Guard - Create 3 crystals around the player. Each crystal damages nearby enemies and shatters in a small freezing blast.", "List Paragraph"),
        ("Lv2: Reinforced Ice - Create 4 crystals with a larger damage radius and longer duration.", "List Paragraph"),
        ("Lv3: Splinter Burst - Shattered crystals fire 3 small splinters toward nearby enemies.", "List Paragraph"),
        ("Lv4: Permafrost - Crystals continuously slow nearby enemies, and the shattering blast freezes them for longer.", "List Paragraph"),
        ("Lv5: Crystal Sanctuary - Create 6 crystals in a wider ring. When the final crystal shatters, all remaining crystals detonate together in a powerful frost nova.", "List Paragraph"),
    ],
)

replace_section(
    doc,
    "Ability 3: Stone Spikes",
    "Ability 4:RockFall",
    [
        ("Description", "Normal"),
        ("Stone spikes erupt from the ground at several positions around the player, prioritizing nearby enemy groups. Each eruption deals heavy damage, briefly pins normal enemies in place, and leaves cracked ground that damages enemies standing on it.", "Normal"),
        ("Combat Role", "Normal"),
        ("A targeted area-control ability that interrupts clustered enemies and creates short-lived zones of dangerous ground.", "Normal"),
        ("Upgrades", "Normal"),
        ("Lv1: Earthen Eruption - Summon 3 spikes near nearby enemies. Each spike damages and briefly pins normal enemies.", "List Paragraph"),
        ("Lv2: Jagged Field - Summon 5 spikes with a larger impact area. Impaled enemies begin bleeding.", "List Paragraph"),
        ("Lv3: Fault Line - Each spike sends a crack toward another nearby enemy, causing a smaller follow-up eruption.", "List Paragraph"),
        ("Lv4: Seismic Rhythm - Eruptions occur in two waves, and cracked ground remains briefly to deal damage over time.", "List Paragraph"),
        ("Lv5: Worldspine - Summon a ring of massive spikes followed by a central eruption. The final eruption launches enemies outward and deals bonus damage to elites and bosses.", "List Paragraph"),
    ],
)

doc.save(OUTPUT)
print(OUTPUT)
